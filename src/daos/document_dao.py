import docx
from models.document import Document


class DocumentDao:
    def __init__(self, file_path):
        self.document = Document(file_path)

        self.initialise_document()

    def initialise_document(self):
        self.document.tables = self.document.doc.tables
        self.document.paras = list(self.set_paragraphs())
        self.document.words = list(self.set_words())

    def set_paragraphs(self):
        counter = 0

        while self.document.doc.paragraphs[counter].text != 'Appendices':
            paragraph = self.document.doc.paragraphs[counter]
            counter += 1

            if not paragraph.style.name.startswith('Heading') and \
                paragraph.text != '' and not paragraph.text.startswith('Figure'):
                yield paragraph

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text != '':
                            yield paragraph

    def set_words(self):
        for paragraph in self.document.paras:
            paragraph = paragraph.text
            words = paragraph.split()
            for word in words:
                if word != ('.' or '–'):
                    yield word

    def print_word_count(self):
        print(len(self.document.words))
